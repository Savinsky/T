# -*- coding: utf-8 -*-
from docxtpl import DocxTemplate
import jinja2
import csv
import docx
import json

#1
doc = DocxTemplate("my_doc.docx")
context = { 'brand': 'brand',
            'model': 'model',
            'consumption': 'consumption',
            'cost': 'cost'
            }
doc.rander(context)
doc.save("generated.docx")

#2
def get_text():
    doc = docx.Document("my_doc.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText
with open('example.csv', 'wb') as f:
    writer = csv.writer(f, delimiter = '\n')
    writer.writerow(get_text())

#3
#t = docx.Document("my_doc.docx")
dict_ex = {'brand': 'Volvo', 'Price': '1.5', 'Vol': 2.0}
dict_to_json = json.dumps(dict_ex)
print(type(dict_to_json), dict_to_json)

with open('dict_to_json.txt', 'w') as f:
    json.dump(dict_ex, f)
